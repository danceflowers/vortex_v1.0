#!/usr/bin/env python3
"""Generate TensorCore Control Module RTL Spec (docx)."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Style setup ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '微软雅黑'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    hs.font.color.rgb = RGBColor(0, 0, 0)


def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return t


def add_code(text):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    return p


# ╔══════════════════════════════════════════════════════════════════╗
# ║                        Title Page                               ║
# ╚══════════════════════════════════════════════════════════════════╝
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TensorCore 控制模块 RTL Spec')
run.font.size = Pt(22)
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('TC Decode · TC Scheduler · Instr Issue · Retire')
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Version 1.0  |  2026-04-13')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# ╔══════════════════════════════════════════════════════════════════╗
# ║                  Module 1: TC Decode                            ║
# ╚══════════════════════════════════════════════════════════════════╝
doc.add_heading('1  TC Decode Module', level=1)

# ── 1.1 ──
doc.add_heading('1.1  功能定位', level=2)
doc.add_paragraph(
    'TC Decode 是 TensorCore 协处理器内部的指令解码模块。'
    'Vortex Core 的 decode 阶段仅识别"这是一条 tensor 指令"，'
    '随后将完整的 32 位原始指令编码原样转发给 TensorCore。'
    'Core 不解析 tensor 指令的任何内部字段。'
)
doc.add_paragraph('TC Decode 接收完整指令字后，负责：')
doc.add_paragraph('1) 字段提取：从 32 位 R-Type 指令中提取 opcode、funct3、funct7、rd、rs1、rs2 各字段，'
                  '进一步解析出指令类型、目标操作数、精度、控制位等。', style='List Number')
doc.add_paragraph('2) 合法性检查：校验指令编码合法性（非法 opcode、保留位非零等），输出 dec_illegal 信号。', style='List Number')
doc.add_paragraph('3) 语义展开：将提取的字段组合为内部控制信号，供 Instr Issue Module 和 TC Scheduler 消费。', style='List Number')

# ── 1.2 ──
doc.add_heading('1.2  设计核心：精度在边界，计算无精度', level=2)
doc.add_paragraph(
    'TensorCore 内部存储采用固定格式：AMem/BMem 存储 fp9（E5M3），CMem/DMem 存储 fp22（E8M14）。'
)
doc.add_paragraph(
    '精度转换仅发生在 MMA_LOAD 和 MMA_STORE 的边界：'
)
doc.add_paragraph('MMA_LOAD 将 TMEM 中的任意精度数据（fp32/fp16/fp8/fp4）转换为 fp9（A/B）或 fp22（C）写入内部 SRAM。',
                  style='List Bullet')
doc.add_paragraph('MMA_STORE 将 DMem 中的 fp22 数据转换为指定精度写回 TMEM。', style='List Bullet')
doc.add_paragraph('WMMA 始终执行 fp9×fp9→fp22 累加，不需要精度参数。', style='List Bullet')
doc.add_paragraph(
    '因此，精度和稀疏模式编码在 MMA_LOAD/MMA_STORE 指令中。'
    '每条指令自包含，TC Decode 为纯组合逻辑无状态模块。'
)

# ── 1.3 ──
doc.add_heading('1.3  指令编码概要', level=2)
doc.add_paragraph('所有 tensor 指令采用标准 RISC-V 32 位 R-Type 格式：')
add_code('[31:25] funct7 (7)  [24:20] rs2 (5)  [19:15] rs1 (5)\n'
         '[14:12] funct3 (3)  [11:7]  rd  (5)  [6:0]   opcode (7)')

add_table(
    ['指令', 'opcode', 'funct3', 'funct7', 'rd', 'rs1', 'rs2'],
    [
        ['MMA_LOAD',  '0001011', '000', 'FMT[1:0] SPARSE[1:0] NUM[2:0]', 'TF dst', 'taddr', '0'],
        ['MMA_STORE', '0101011', '000', 'FMT[1:0] reserved[4:0]',        '0',      'taddr', 'TF src'],
        ['WMMA',      '1011011', '000', 'reserved[3:0] TAIL OR WS',      'ctx_id D_taddr', 'A_taddr', '—'],
    ],
)

doc.add_paragraph()
doc.add_paragraph('各字段编码：')
add_table(
    ['字段', '位宽', '编码'],
    [
        ['FMT[1:0]',    '2', '00=fp32, 01=fp16, 10=fp8, 11=fp4(保留)'],
        ['SPARSE[1:0]', '2', '00=dense, 01=2:4, 10=1:4, 11=保留'],
        ['OR',           '1', '0=非驻留（C 从 CMem passthrough），1=驻留（accum_fifo 循环累加）'],
        ['WS',           '1', '0=正常（每轮 fill B），1=weight-stationary（BMem 锁定）'],
        ['TAIL',         '1', '0=非尾部，1=驻留累加序列的最后一条 WMMA（OR=0 时忽略）'],
    ],
)

# ── 1.4 ──
doc.add_heading('1.4  在协处理器中的位置', level=2)
add_code(
    'Vortex Core\n'
    '    │  Decode：识别 "tensor 指令"\n'
    '    │  Execute：原样转发 32b + warp_id\n'
    '    ▼\n'
    '┌──────────────────────────────┐\n'
    '│         TC Decode             │\n'
    '│  (纯组合逻辑，无状态)         │\n'
    '│                               │\n'
    '│  ┌────────────────┐          │\n'
    '│  │ otc_instr_split│ 切字段   │\n'
    '│  └───────┬────────┘          │\n'
    '│          ▼                    │\n'
    '│  ┌────────────────┐          │\n'
    '│  │  otc_decoder   │          │\n'
    '│  │ 识别+展开+合法 │          │\n'
    '│  └───────┬────────┘          │\n'
    '│          ▼                    │\n'
    '│    dec_* 输出信号             │\n'
    '└──────┬───────────────────────┘\n'
    '       │\n'
    '       ├──→ Instr Issue Module\n'
    '       └──→ TC Scheduler'
)

# ── 1.5 ──
doc.add_heading('1.5  设计原则', level=2)
doc.add_paragraph('Core 零解释：Vortex Core 对 tensor 指令只做识别和转发，不提取任何内部字段。'
                  'TensorCore 作为协处理器可移植到其他 GPU 核心，仅需约定 32 位指令转发接口。', style='List Bullet')
doc.add_paragraph('两级解码：第一级 otc_instr_split 纯组合位域切割；'
                  '第二级 otc_decoder 负责指令识别、参数展开、合法性检查。', style='List Bullet')
doc.add_paragraph('解码零延迟：全组合逻辑，无内部状态寄存器，同一周期输出解码结果。', style='List Bullet')
doc.add_paragraph('无独立 AC 位：每条 WMMA 始终处理 C 操作数。OR=0 时 C 从 CMem 读取（passthrough），'
                  'OR=1 时 C 从 accum_fifo 取（循环累加）。需要 D=A×B 效果时，软件往 CMem 加载零值。', style='List Bullet')

# ── 1.6 ──
doc.add_heading('1.6  四种 WMMA 工作模式', level=2)
add_table(
    ['WS', 'OR', 'AMem', 'BMem', 'C 来源', 'D 去向'],
    [
        ['0', '0', '每轮 fill', '每轮 fill', 'CMem passthrough', 'DMem'],
        ['0', '1', '每轮 fill', '每轮 fill', 'accum_fifo 循环', 'FIFO → 尾部 job → DMem'],
        ['1', '0', '每轮 fill', '首次 fill 后锁定', 'CMem passthrough', 'DMem'],
        ['1', '1', '每轮 fill', '首次 fill 后锁定', 'accum_fifo 循环', 'FIFO → 尾部 job → DMem'],
    ],
)

# ── Ch2: TC Decode 接口定义 ──
doc.add_heading('1.7  接口定义', level=2)

doc.add_heading('1.7.1  时钟与复位', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['input', 'clk', '1', '时钟'],
    ['input', 'reset_n', '1', '异步低有效复位'],
])

doc.add_heading('1.7.2  Core 转发接口', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['input', 'core_dispatch_valid', '1', 'Core 本周期转发了一条 tensor 指令'],
    ['input', 'core_warp_id', 'LOG2(NUM_WARPS)', '发射 warp 的 ID'],
    ['input', 'core_raw_instr', '32', '完整 R-Type tensor 指令编码'],
])

doc.add_heading('1.7.3  通用解码输出', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'dec_valid', '1', '本周期有有效解码结果'],
    ['output', 'dec_is_tcu', '1', '确认为合法 tensor 指令'],
    ['output', 'dec_illegal', '1', '非法指令'],
])

doc.add_heading('1.7.4  指令类别输出', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'dec_is_mma_load', '1', 'MMA_LOAD'],
    ['output', 'dec_is_mma_store', '1', 'MMA_STORE'],
    ['output', 'dec_is_wmma', '1', 'WMMA'],
])

doc.add_heading('1.7.5  寄存器字段输出（直接切割）', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'dec_rd', '5', 'rd 字段原始值'],
    ['output', 'dec_rs1', '5', 'rs1 字段原始值'],
    ['output', 'dec_rs2', '5', 'rs2 字段原始值'],
])

doc.add_heading('1.7.6  MMA_LOAD / MMA_STORE 专用字段', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'dec_fmt', '2', 'FMT[1:0]：外部精度（00=fp32, 01=fp16, 10=fp8, 11=fp4）'],
    ['output', 'dec_sparse', '2', 'SPARSE[1:0]：稀疏模式（00=dense, 01=2:4, 10=1:4）。仅 MMA_LOAD 有效'],
    ['output', 'dec_num', '3', 'NUM[2:0]：搬运数量'],
])

doc.add_heading('1.7.7  WMMA 专用字段', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'dec_or', '1', 'output_resident（FIFO 驻留累加）'],
    ['output', 'dec_ws', '1', 'weight-stationary（BMem 锁定）'],
    ['output', 'dec_tail', '1', 'TAIL（驻留序列尾部标记，OR=0 时忽略）'],
])

doc.add_heading('1.7.8  内部子模块结构', level=3)
doc.add_paragraph(
    'otc_instr_split（第一级）：纯组合逻辑，将 32 位指令按 R-Type 格式切割为 7 个字段。无状态，无时序。'
)
doc.add_paragraph(
    'otc_decoder（第二级）：组合逻辑，根据 opcode+funct3 判断指令类型，从 funct7 提取控制字段，执行合法性检查。'
)

doc.add_heading('1.7.9  合法性检查规则', level=3)
add_table(['条件', '结果'], [
    ['opcode 不在 {0001011, 0101011, 1011011} 中', 'dec_illegal = 1'],
    ['MMA_LOAD 的 SPARSE ≠ 00 但目标非 A/B', 'dec_illegal = 1'],
    ['funct7 保留位非零', 'dec_illegal = 1'],
    ['FMT = 11（fp4 保留）', 'dec_illegal = 1（当前版本）'],
])

doc.add_page_break()

# ╔══════════════════════════════════════════════════════════════════╗
# ║                  Module 2: TC Scheduler                         ║
# ╚══════════════════════════════════════════════════════════════════╝
doc.add_heading('2  TC Scheduler Module', level=1)

doc.add_heading('2.1  功能定位', level=2)
doc.add_paragraph(
    'TC Scheduler 是 TensorCore 协处理器内部的指令发射调度模块。'
    '它位于 Vortex Core 的 Warp Scheduler 与 TensorCore 内部资源之间，负责：'
)
doc.add_paragraph('1) 发射评分：为 Warp Scheduler 提供 tensor 指令的优先级评分（scheduler_score），'
                  'Warp Scheduler 据此在多个 warp 的候选指令中选择优先发射的 warp。', style='List Number')
doc.add_paragraph('2) 发射条件检查：根据内部 Slot 状态和 TMEM Handle 状态，判断当前 tensor 指令能否发射。', style='List Number')

doc.add_heading('2.2  评分体系', level=2)
add_table(
    ['评分', '含义', '条件'],
    [
        ['0', '不可发射', '前置条件未满足'],
        ['3', 'MMA_STORE 就绪', 'DMem 数据有效、Store 通道空闲'],
        ['4', 'MMA_LOAD 就绪', 'TMEM Handle 就绪、目标 Slot 可接受新 fill'],
        ['5', 'WMMA 就绪（最高优先级）', 'A/B valid、C valid、无 pending'],
    ],
)
doc.add_paragraph('WMMA 优先级最高：一旦操作数就绪应尽快发射，以保持计算阵列利用率。')

doc.add_heading('2.3  Slot 状态观测机制', level=2)
doc.add_paragraph(
    'TC Scheduler 读取上一周期末的快照（published slot state），对应 RTL 中的寄存器采样。'
    '1 周期延迟避免 Scheduler 评分路径上出现跨模块组合逻辑长链。'
)

doc.add_heading('2.4  设计原则', level=2)
doc.add_paragraph('只读观测：TC Scheduler 不修改任何 Slot 或 Handle 状态，只做查询和评分。', style='List Bullet')
doc.add_paragraph('单周期响应：评分为组合逻辑，Warp Scheduler 同周期可得到结果。', style='List Bullet')
doc.add_paragraph('与 Instr Issue 解耦：Scheduler 负责"能不能发"，Instr Issue 负责"怎么发"。', style='List Bullet')

doc.add_heading('2.5  接口定义', level=2)

doc.add_heading('2.5.1  Warp Scheduler 查询接口', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['input', 'query_valid', '1', 'Warp Scheduler 发起评分查询'],
    ['input', 'query_warp_id', 'LOG2(NUM_WARPS)', '查询的 warp ID'],
    ['input', 'query_type', '2', '指令类型（MMA_LOAD=01, MMA_STORE=10, WMMA=11）'],
    ['output', 'score', '3', '发射优先级（0/3/4/5）'],
])

doc.add_heading('2.5.2  Published Slot 状态输入', level=3)
doc.add_paragraph('A Slot：')
add_table(['方向', '信号名', '位宽', '描述'], [
    ['input', 'pub_a_ready', '1', '可接受新 fill'],
    ['input', 'pub_a_pending', '1', 'fill 进行中'],
    ['input', 'pub_a_valid', '1', 'AMem 数据有效（完整 m16n16k16）'],
    ['input', 'pub_a_wmma_pending', '1', '被 WMMA 引用中'],
])

doc.add_paragraph('B Slot：')
add_table(['方向', '信号名', '位宽', '描述'], [
    ['input', 'pub_b_ready', '1', '可接受新 fill'],
    ['input', 'pub_b_pending', '1', 'fill 进行中'],
    ['input', 'pub_b_valid', '1', 'BMem 数据有效（完整 m16n16k16）'],
    ['input', 'pub_b_wmma_pending', '1', '被 WMMA 引用中'],
    ['input', 'pub_b_ws_locked', '1', 'weight-stationary 锁定'],
])

doc.add_paragraph('C Slot：')
add_table(['方向', '信号名', '位宽', '描述'], [
    ['input', 'pub_c_ready', '1', '可接受新 fill'],
    ['input', 'pub_c_pending', '1', 'C bias fill 进行中'],
    ['input', 'pub_c_valid', '1', 'CMem 数据有效'],
    ['input', 'pub_c_wmma_inflight', 'N', '绑定的未完成 WMMA 数'],
])

doc.add_paragraph('DMem：')
add_table(['方向', '信号名', '位宽', '描述'], [
    ['input', 'pub_dmem_ready', '1', '可接受新 WMMA 输出'],
    ['input', 'pub_dmem_pending', '1', 'Retire Unit 写入中'],
    ['input', 'pub_dmem_valid', '1', '全部 subtile 写入完成'],
    ['input', 'pub_dmem_store_pending', '1', 'MMA_STORE 进行中（DMem→TMEM）'],
])

doc.add_heading('2.5.3  TMEM Handle 状态输入', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['input', 'handle_payload_ready', '1', 'payload 数据已就绪'],
    ['input', 'handle_meta_ready', '1', 'metadata 已就绪'],
    ['input', 'handle_busy', '1', '正在被 TMA 操作占用'],
])

doc.add_heading('2.5.4  评分逻辑', level=3)
add_code(
    'WMMA (score=5 or 0):\n'
    '    pub_a_valid && !pub_a_wmma_pending\n'
    '    && pub_b_valid && !pub_b_wmma_pending\n'
    '    && pub_c_valid\n'
    '\n'
    'MMA_LOAD (score=4 or 0):\n'
    '    handle_payload_ready && !handle_busy\n'
    '    && 目标 slot ready=1\n'
    '\n'
    'MMA_STORE (score=3 or 0):\n'
    '    !handle_busy\n'
    '    && pub_dmem_valid && !pub_dmem_store_pending'
)

doc.add_page_break()

# ╔══════════════════════════════════════════════════════════════════╗
# ║                  Module 3: Instr Issue Module                   ║
# ╚══════════════════════════════════════════════════════════════════╝
doc.add_heading('3  Instr Issue Module', level=1)

doc.add_heading('3.1  功能定位', level=2)
doc.add_paragraph(
    'Instr Issue Module 是 TensorCore 协处理器的核心指令执行控制器，'
    '负责三类 tensor 指令的执行控制和内部 SRAM slot 状态管理。'
    '本模块是纯控制模块，不经手数据通路。'
)

add_table(
    ['指令', '执行动作'],
    [
        ['MMA_LOAD', '接收解码结果 → 更新 slot 状态为 PENDING → 向 TMEM 读端口发出 fill 请求 → fill 完成后更新 slot 为 VALID'],
        ['WMMA', '从 job 寄存器取活跃 job → 按 uop 计数器逐周期发射 8 个 8×8 原语到计算阵列 → 全部发射后释放 slot'],
        ['MMA_STORE', '检查 DMem VALID → 向 TMEM 写端口发出 store 请求 → store 完成后释放 DMem'],
    ],
)

doc.add_heading('3.2  WMMA 原语分解与 SRAM 寻址', level=2)
doc.add_paragraph('一条 WMMA（m16n16k16）分解为 8 个 8×8 原语。地址计算公式：')
add_code(
    'amem_line = k_phase × 2 + (uop_counter % 4) / 2\n'
    'bmem_line = k_phase × 2 + (uop_counter % 4) % 2\n'
    'cmem_line = uop_counter % 4\n'
    'k_phase   = uop_counter / 4'
)

doc.add_paragraph()
add_table(
    ['uop_id', 'amem_line', 'bmem_line', 'cmem_line', '说明'],
    [
        ['0', '0', '0', '0', 'K-phase 0, M-block 0 × N-block 0'],
        ['1', '0', '1', '1', 'K-phase 0, M-block 0 × N-block 1'],
        ['2', '1', '0', '2', 'K-phase 0, M-block 1 × N-block 0'],
        ['3', '1', '1', '3', 'K-phase 0, M-block 1 × N-block 1'],
        ['4', '2', '2', '0', 'K-phase 1, M-block 0 × N-block 0'],
        ['5', '2', '3', '1', 'K-phase 1, M-block 0 × N-block 1'],
        ['6', '3', '2', '2', 'K-phase 1, M-block 1 × N-block 0'],
        ['7', '3', '3', '3', 'K-phase 1, M-block 1 × N-block 1'],
    ],
)

doc.add_paragraph(
    'AMem/BMem 各 4 line（line 0-1 = K-phase 0, line 2-3 = K-phase 1），CMem 4 line（4 个 subtile 循环访问）。'
    'K-phase 0 和 K-phase 1 的结果累加到同一个 cmem_line 位置。'
)

doc.add_heading('3.3  16×16 输出 tile 的 subtile 布局', level=2)
add_code(
    '         col 0-7      col 8-15\n'
    '       ┌──────────┬──────────┐\n'
    'row    │subtile 0 │subtile 1 │\n'
    '0-7    │(m=0,n=0) │(m=0,n=1) │\n'
    '       ├──────────┼──────────┤\n'
    'row    │subtile 2 │subtile 3 │\n'
    '8-15   │(m=1,n=0) │(m=1,n=1) │\n'
    '       └──────────┴──────────┘'
)

doc.add_heading('3.4  Slot 状态机', level=2)
doc.add_paragraph('所有 SRAM 遵循统一的三态模式：READY → PENDING → VALID → READY。')
add_table(
    ['SRAM', 'READY→PENDING', 'PENDING→VALID', 'VALID→READY'],
    [
        ['AMem', 'MMA_LOAD(A) 接收', '全部 packet 写入完成', '8 个 uop 全部 push 后释放'],
        ['BMem', 'MMA_LOAD(B) 接收', '全部 packet 写入完成', 'ws=0: 同 AMem；ws=1: 保持 VALID + ws_locked'],
        ['CMem', 'MMA_LOAD(C) 接收', '全部 subtile 写入完成', '8 个 uop 全部 push 后释放'],
        ['DMem', 'WMMA 指令接收时分配', 'Retire Module 全部 subtile 写入完成', 'MMA_STORE 全部 packet 写回完成后释放'],
    ],
)

doc.add_heading('3.5  指令执行 FSM', level=2)
add_table(
    ['状态', '含义', '退出条件'],
    [
        ['IDLE', '等待新指令', 'dec_valid 且对应指令类型有效'],
        ['FILL_ACTIVE', 'MMA_LOAD 接收 TMEM packet 中', 'tmem_rd_done=1'],
        ['WMMA_ISSUING', 'WMMA 逐周期发射原语', 'uop_counter==8'],
        ['STORE_ACTIVE', 'MMA_STORE 逐 subtile 写回 TMEM', 'tmem_wr_done=1'],
    ],
)
doc.add_paragraph('三路互斥：同一时刻 FSM 只处于一种活跃状态，不并行执行多种操作。')

doc.add_heading('3.6  MMA_LOAD 执行流程', level=2)
doc.add_paragraph('① 接收 TC Decode 输出 → 目标 slot 检查（必须 READY）→ slot 状态更新为 PENDING。')
doc.add_paragraph('② 向 TMEM 读端口发出请求（基于 handle + tile_id）。')
doc.add_paragraph('③ 逐 packet 接收 TMEM 响应，驱动 SRAM 写控制信号（地址递增）。'
                  '精度转换由 SRAM 前端独立转换模块完成。')
doc.add_paragraph('④ 全部 packet 接收完毕 → slot 状态更新为 VALID → FSM 回到 IDLE。')

doc.add_paragraph('Packet 数量：')
add_table(
    ['目标', 'fp16', 'fp8', '说明'],
    [
        ['AMem (m16×k16)', '8 packets', '4 packets', '4 line × 2 packets (fp16) 或 1 packet (fp8)'],
        ['BMem (k16×n16)', '8 packets', '4 packets', '同 AMem'],
        ['CMem (m16×n16)', 'fp32:16p, fp16:8p, fp8:4p', '—', '4 subtile × packets_per_subtile'],
    ],
)

doc.add_heading('3.7  WMMA 执行流程', level=2)
doc.add_paragraph('① 接收 TC Decode 输出 → 前置条件检查（A valid, B valid, C valid）→ 创建活跃 job，'
                  '锁存 wgid/async_id/OR/WS/TAIL，uop_counter=0。')
doc.add_paragraph('② Slot 标记：A/B wmma_pending=1，C wmma_inflight++，DMem READY→PENDING。')
doc.add_paragraph('③ 计算 materialize = (OR==0) || (TAIL==1)。')
doc.add_paragraph('④ 逐周期发射循环（最快 8 周期，每周期 1 个原语）：')
doc.add_paragraph('    a) 根据 uop_counter 计算 amem_line, bmem_line, cmem_line', style='List Bullet')
doc.add_paragraph('    b) 检查计算阵列 tc_ready 信号，若 ready=0 则 stall', style='List Bullet')
doc.add_paragraph('    c) 驱动 SRAM 读控制（amem_rd_en/line, bmem_rd_en/line, cmem_rd_en/line）', style='List Bullet')
doc.add_paragraph('    d) 驱动 tc_push_valid=1，输出 meta{wgid, async_id, subtile_id, materialize}', style='List Bullet')
doc.add_paragraph('    e) uop_counter++', style='List Bullet')
doc.add_paragraph('⑤ uop_counter==8 → Slot 释放（A/C: VALID→READY；B: ws=0 释放，ws=1 保持 VALID+ws_locked）→ FSM 回到 IDLE。')

doc.add_heading('3.8  MMA_STORE 执行流程', level=2)
doc.add_paragraph('① 接收 TC Decode 输出 → 前置条件检查（DMem VALID, store_pending=0）→ store_pending=1。')
doc.add_paragraph('② 向 TMEM 写端口发出请求。')
doc.add_paragraph('③ 逐 subtile 从 DMem 读取，驱动精度转换和 TMEM 写控制信号。')
doc.add_paragraph('④ 全部写入完成 → store_pending=0, DMem VALID→READY → FSM 回到 IDLE。')

doc.add_heading('3.9  设计原则', level=2)
doc.add_paragraph('纯控制模块：只生成控制信号（SRAM 读/写地址、TMEM 读/写请求、计算阵列 push 使能等），'
                  '不经手数据通路。精度转换由 SRAM 前端的独立转换模块完成。', style='List Bullet')
doc.add_paragraph('Slot 状态集中管理：所有 SRAM slot 状态寄存器由本模块持有和更新，'
                  'TC Scheduler 通过周期末快照只读访问。', style='List Bullet')
doc.add_paragraph('三路互斥：同一时刻 FSM 只处于一种活跃状态。', style='List Bullet')
doc.add_paragraph('顺序发射：WMMA 的 8 个原语严格按 uop_id 递增顺序发射。', style='List Bullet')
doc.add_paragraph('ws 保持：weight-stationary 模式下 BMem 完成 WMMA 后不释放，保持 VALID + ws_locked，'
                  '直到接收到 ws=0 的 WMMA 指令后在下一次 WMMA 完成时释放。', style='List Bullet')

doc.add_heading('3.10  接口定义', level=2)

doc.add_heading('3.10.1  TC Decode 输入', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['input', 'dec_valid', '1', '有效解码结果'],
    ['input', 'dec_is_mma_load', '1', 'MMA_LOAD 指令'],
    ['input', 'dec_is_mma_store', '1', 'MMA_STORE 指令'],
    ['input', 'dec_is_wmma', '1', 'WMMA 指令'],
    ['input', 'dec_warp_id', 'LOG2(NUM_WARPS)', 'Warp ID'],
    ['input', 'dec_fmt', '2', '外部精度（MMA_LOAD/MMA_STORE 用）'],
    ['input', 'dec_sparse', '2', '稀疏模式（MMA_LOAD 用）'],
    ['input', 'dec_or', '1', 'output_resident'],
    ['input', 'dec_ws', '1', 'weight-stationary'],
    ['input', 'dec_tail', '1', 'TAIL 标记'],
    ['input', 'dec_target', '2', 'MMA_LOAD 目标（A/B/C）'],
    ['input', 'dec_handle', 'H', 'TMEM 分配句柄'],
    ['input', 'dec_tile_id', 'T', 'Tile ID'],
    ['output', 'issue_ready', '1', '本模块可接受新指令（FSM 处于 IDLE）'],
])

doc.add_heading('3.10.2  TMEM 读端口控制（MMA_LOAD）', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'tmem_rd_req', '1', '读请求有效'],
    ['output', 'tmem_rd_handle', 'H', 'TMEM 句柄'],
    ['output', 'tmem_rd_tile_id', 'T', 'Tile ID'],
    ['output', 'tmem_rd_fmt', '2', '源数据精度'],
    ['input', 'tmem_rd_grant', '1', '读请求被接受'],
    ['input', 'tmem_rd_packet_valid', '1', '本周期有 packet 返回'],
    ['input', 'tmem_rd_done', '1', '全部 packet 传输完成'],
])

doc.add_heading('3.10.3  TMEM 写端口控制（MMA_STORE）', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'tmem_wr_req', '1', '写请求有效'],
    ['output', 'tmem_wr_handle', 'H', 'TMEM 句柄'],
    ['output', 'tmem_wr_tile_id', 'T', 'Tile ID'],
    ['output', 'tmem_wr_fmt', '2', '目标数据精度'],
    ['input', 'tmem_wr_grant', '1', '写请求被接受'],
    ['input', 'tmem_wr_done', '1', '全部 packet 写入完成'],
])

doc.add_heading('3.10.4  SRAM 写控制（MMA_LOAD fill）', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'sram_fill_en', '1', 'fill 写使能'],
    ['output', 'sram_fill_target', '2', '目标 SRAM（A/B/C）'],
    ['output', 'sram_fill_line_addr', '2', '目标 line 地址（0-3）'],
    ['output', 'sram_fill_fmt', '2', '源精度（传递给前端精度转换模块）'],
    ['output', 'sram_fill_sparse', '2', '稀疏模式（传递给前端）'],
])

doc.add_heading('3.10.5  SRAM 读控制（WMMA）', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'amem_rd_en', '1', 'AMem 读使能'],
    ['output', 'amem_rd_line', '2', 'AMem line 地址（0-3）'],
    ['output', 'bmem_rd_en', '1', 'BMem 读使能'],
    ['output', 'bmem_rd_line', '2', 'BMem line 地址（0-3）'],
    ['output', 'cmem_rd_en', '1', 'CMem 读使能（OR=0 时有效）'],
    ['output', 'cmem_rd_line', '2', 'CMem line 地址（0-3）'],
])

doc.add_heading('3.10.6  DMem 读控制（MMA_STORE）', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'dmem_rd_en', '1', 'DMem 读使能'],
    ['output', 'dmem_rd_subtile', '2', 'subtile 地址（0-3）'],
    ['output', 'dmem_rd_fmt', '2', '目标精度'],
])

doc.add_heading('3.10.7  计算阵列接口', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'tc_push_valid', '1', '本周期发射一个原语'],
    ['output', 'tc_push_wgid', 'W', 'workgroup ID'],
    ['output', 'tc_push_async_id', 'A', '异步操作 ID'],
    ['output', 'tc_push_subtile_id', '2', 'subtile 编号（0-3）'],
    ['output', 'tc_push_materialize', '1', '1=结果写 DMem，0=留在 FIFO'],
    ['output', 'tc_push_or', '1', 'output_resident 模式'],
    ['input', 'tc_ready', '1', '计算阵列可接受输入'],
])

doc.add_heading('3.10.8  Retire Module 交互', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['input', 'retire_wmma_done', '1', '一个 WMMA 全部退休'],
    ['input', 'retire_wmma_async_id', 'A', '对应 async_id'],
    ['input', 'retire_dmem_valid', '1', 'DMem 全部 subtile 写入完成'],
])

doc.add_heading('3.10.9  Slot 状态快照输出（→ TC Scheduler）', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'pub_a_ready', '1', 'A slot 可接受 fill'],
    ['output', 'pub_a_pending', '1', 'A slot fill 中'],
    ['output', 'pub_a_valid', '1', 'A slot 数据有效'],
    ['output', 'pub_a_wmma_pending', '1', 'A slot 被 WMMA 引用'],
    ['output', 'pub_b_ready', '1', 'B slot 可接受 fill'],
    ['output', 'pub_b_pending', '1', 'B slot fill 中'],
    ['output', 'pub_b_valid', '1', 'B slot 数据有效'],
    ['output', 'pub_b_wmma_pending', '1', 'B slot 被 WMMA 引用'],
    ['output', 'pub_b_ws_locked', '1', 'B slot ws 锁定'],
    ['output', 'pub_c_ready', '1', 'C slot 可接受 fill'],
    ['output', 'pub_c_pending', '1', 'C slot fill 中'],
    ['output', 'pub_c_valid', '1', 'C slot 数据有效'],
    ['output', 'pub_c_wmma_inflight', 'N', 'C slot 未完成 WMMA 数'],
    ['output', 'pub_dmem_ready', '1', 'DMem 可接受写入'],
    ['output', 'pub_dmem_pending', '1', 'DMem 写入中'],
    ['output', 'pub_dmem_valid', '1', 'DMem 数据有效'],
    ['output', 'pub_dmem_store_pending', '1', 'MMA_STORE 中'],
])

doc.add_page_break()

# ╔══════════════════════════════════════════════════════════════════╗
# ║                  Module 4: Retire Module                        ║
# ╚══════════════════════════════════════════════════════════════════╝
doc.add_heading('4  Retire Module', level=1)

doc.add_heading('4.1  功能定位', level=2)
doc.add_paragraph(
    'Retire Module 是 TensorCore 协处理器的计算结果接收与写回控制器。'
    '它位于计算阵列输出端，负责：'
)
doc.add_paragraph('1) 接收退休原语：从计算阵列输出端逐周期接收退休的 8×8 fp22 结果及其 meta。', style='List Number')
doc.add_paragraph('2) 条件写入 DMem：根据 meta 中的 materialize 标志决定是否写入 DMem。', style='List Number')
doc.add_paragraph('3) 多 WMMA 完成追踪：通过 pending_table 按 async_id 追踪多个在飞 WMMA 的退休进度，'
                  '全部退休后通知 Instr Issue Module。', style='List Number')

doc.add_heading('4.2  退休流程', level=2)
doc.add_paragraph('每周期（retire_valid=1 时）执行以下步骤：')
add_table(
    ['步骤', '动作'],
    [
        ['① 查表', '在 pending_table 中查找 retire_async_id 匹配的 entry'],
        ['② 递减计数', '匹配 entry 的 count--'],
        ['③ 条件写 DMem', 'materialize=1 时：dmem_wr_en=1；materialize=0 时：不写'],
        ['④ 更新位图', 'materialize=1 时：subtile_written[subtile_id]=1'],
        ['⑤ 完成检查', 'count==0 → 触发 notify_wmma_done，释放 entry'],
        ['⑥ DMem 有效检查', "subtile_written==4'b1111 → 触发 notify_dmem_valid"],
    ],
)

doc.add_heading('4.3  多 WMMA 在飞支持', level=2)
doc.add_paragraph(
    '计算阵列管线延迟 ~12 周期，WMMA 发射间隔最短 8 周期。管线中可重叠多个 WMMA。'
    'pending_table 深度 D=4，覆盖极端情况。'
)
add_table(
    ['字段', '位宽', '描述'],
    [
        ['valid', '1', 'entry 有效'],
        ['async_id', 'A', '追踪的 async_id'],
        ['count', '4', '剩余待退休原语数（初始 8）'],
    ],
)

doc.add_heading('4.4  materialize 标志', level=2)
add_table(
    ['指令 OR', '指令 TAIL', 'materialize', '含义'],
    [
        ['0', 'x（忽略）', '1', '非驻留模式，每个原语结果都写 DMem'],
        ['1', '0', '0', '驻留中间 WMMA，结果留在 FIFO，不写 DMem'],
        ['1', '1', '1', '驻留尾部 WMMA，结果退出 FIFO 写 DMem'],
    ],
)
doc.add_paragraph(
    'materialize 由 Instr Issue Module 在发射时计算，打包进 meta，'
    '随原语流过计算阵列管线，到达 Retire Module 时直接使用。'
)

doc.add_heading('4.5  计算阵列对 materialize 的响应', level=2)
add_table(
    ['materialize', '计算阵列 final_add 行为'],
    [
        ['0', '结果输出到 Retire Module 且回灌 accum_fifo（循环累加继续）'],
        ['1', '结果输出到 Retire Module，不回灌 accum_fifo（累加序列结束）'],
    ],
)

doc.add_heading('4.6  设计原则', level=2)
doc.add_paragraph('被动接收：不主动拉取数据，只在 retire_valid 有效时响应。', style='List Bullet')
doc.add_paragraph('多 WMMA 追踪：pending_table 按 async_id 索引，支持管线中多个 WMMA 重叠。', style='List Bullet')
doc.add_paragraph('materialize 驱动写入：DMem 写入完全由 materialize 标志控制，'
                  'Retire Module 无需了解 OR/TAIL 语义。', style='List Bullet')
doc.add_paragraph('与 Instr Issue 分工明确：Instr Issue 管发射端，Retire 管接收端。', style='List Bullet')

doc.add_heading('4.7  接口定义', level=2)

doc.add_heading('4.7.1  计算阵列输出接口', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['input', 'retire_valid', '1', '本周期有原语退休'],
    ['input', 'retire_wgid', 'W', 'workgroup ID'],
    ['input', 'retire_async_id', 'A', '异步操作 ID'],
    ['input', 'retire_subtile_id', '2', 'subtile 编号（0-3）'],
    ['input', 'retire_materialize', '1', '1=写 DMem，0=仅计数'],
    ['input', 'retire_data', '8×8×22', 'fp22 结果（数据通路直连）'],
])

doc.add_heading('4.7.2  DMem 写端口', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'dmem_wr_en', '1', '写使能（= retire_valid & retire_materialize）'],
    ['output', 'dmem_wr_subtile', '2', 'subtile 地址（0-3）'],
    ['output', 'dmem_wr_data', '8×8×22', 'fp22 写数据（数据通路直连）'],
])

doc.add_heading('4.7.3  Instr Issue Module 通知', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['output', 'notify_wmma_done', '1', '一个 WMMA 的 8 个原语全部退休'],
    ['output', 'notify_wmma_async_id', 'A', '完成的 async_id'],
    ['output', 'notify_dmem_valid', '1', 'DMem 全部 subtile 写入完成'],
])

doc.add_heading('4.7.4  Instr Issue Module 配置', level=3)
add_table(['方向', '信号名', '位宽', '描述'], [
    ['input', 'cfg_alloc_valid', '1', '分配新 pending_table entry'],
    ['input', 'cfg_alloc_async_id', 'A', '新 WMMA 的 async_id'],
])

doc.add_heading('4.7.5  内部寄存器', level=3)
doc.add_paragraph('pending_table（深度 D=4）：')
add_table(['字段', '位宽', '描述'], [
    ['valid', '1', 'entry 有效'],
    ['async_id', 'A', '追踪的 async_id'],
    ['count', '4', '剩余待退休原语数（初始 8）'],
])

doc.add_paragraph()
doc.add_paragraph('DMem 追踪：')
add_table(['寄存器', '位宽', '描述'], [
    ['subtile_written', '4', 'subtile 写入位图（bit i = subtile i 已写入）'],
])

# ── Save ──
out_path = '/mnt/d/wode_code_trunk/vortex/spec/tc_control_modules_spec.docx'
doc.save(out_path)
print(f'Saved to {out_path}')
